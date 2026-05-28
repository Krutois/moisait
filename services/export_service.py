from io import BytesIO
import os
import textwrap


EXPORT_LABELS = {
    "en": {
        "title": "Title", "date": "Date", "language": "Language", "source": "Source",
        "words": "Words", "duration": "Duration", "summary": "Study summary",
        "keywords": "Keywords", "tags": "Tags", "transcript": "Transcript",
        "document_title": "SmartLecture Transcript",
    },
    "ru": {
        "title": "Название", "date": "Дата", "language": "Язык", "source": "Источник",
        "words": "Слова", "duration": "Длительность", "summary": "Учебный конспект",
        "keywords": "Ключевые слова", "tags": "Теги", "transcript": "Расшифровка",
        "document_title": "Расшифровка SmartLecture",
    },
    "kk": {
        "title": "Атауы", "date": "Күні", "language": "Тілі", "source": "Дереккөзі",
        "words": "Сөздер", "duration": "Ұзақтығы", "summary": "Оқу конспектісі",
        "keywords": "Кілт сөздер", "tags": "Тегтер", "transcript": "Мәтін",
        "document_title": "SmartLecture мәтіні",
    },
}


def labels(lang="en"):
    return EXPORT_LABELS.get(lang, EXPORT_LABELS["en"])


class ExportService:
    @staticmethod
    def build_text(transcription, lang="en"):
        label = labels(lang)
        lines = [
            "SmartLecture",
            f"{label['title']}: {transcription.title or label['transcript']}",
            f"{label['date']}: {transcription.created_at.strftime('%Y-%m-%d %H:%M') if transcription.created_at else '-'}",
            f"{label['language']}: {transcription.language or '-'}",
            f"{label['source']}: {transcription.source or '-'}",
            f"{label['words']}: {transcription.word_count or 0}",
            f"{label['duration']}: {transcription.duration or 0}",
            "",
        ]

        if transcription.summary:
            lines.extend([f"{label['summary']}:", transcription.summary, ""])

        if transcription.keywords_json:
            lines.extend([f"{label['keywords']}:", ", ".join(transcription.keywords_json), ""])

        if transcription.tags_json:
            lines.extend([f"{label['tags']}:", ", ".join(transcription.tags_json), ""])

        lines.extend([f"{label['transcript']}:", transcription.text or ""])
        return "\n".join(lines)

    @staticmethod
    def txt(transcription, lang="en"):
        return ExportService.build_text(transcription, lang).encode("utf-8")

    @staticmethod
    def _font_path():
        candidates = [
            os.path.join(os.environ.get("WINDIR", "C:\\Windows"), "Fonts", "arial.ttf"),
            os.path.join(os.environ.get("WINDIR", "C:\\Windows"), "Fonts", "calibri.ttf"),
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/dejavu/DejaVuSans.ttf",
            "/Library/Fonts/Arial.ttf",
        ]
        return next((path for path in candidates if os.path.exists(path)), None)

    @staticmethod
    def pdf(transcription, lang="en"):
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.pdfgen import canvas

        font_path = ExportService._font_path()
        if not font_path:
            raise RuntimeError("Unicode font not found. Install Arial or DejaVu Sans.")

        font_name = "SmartLectureUnicode"
        if font_name not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont(font_name, font_path))

        buffer = BytesIO()
        pdf = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        margin = 48
        y = height - margin
        label = labels(lang)

        def draw_line(text="", size=10, color=colors.HexColor("#111827"), leading=15):
            nonlocal y
            if y < margin + 30:
                pdf.showPage()
                y = height - margin
            pdf.setFont(font_name, size)
            pdf.setFillColor(color)
            pdf.drawString(margin, y, text)
            y -= leading

        draw_line("SmartLecture", 20, colors.HexColor("#4f46e5"), 26)
        draw_line(transcription.title or "Transcript", 16, colors.HexColor("#111827"), 24)
        draw_line(
            f"{transcription.created_at.strftime('%Y-%m-%d %H:%M') if transcription.created_at else '-'} | "
            f"{transcription.language or '-'} | {transcription.word_count or 0} {label['words'].lower()}",
            9,
            colors.HexColor("#475569"),
            22,
        )

        for label, value in [
            (label["summary"], transcription.summary),
            (label["keywords"], ", ".join(transcription.keywords_json or [])),
            (label["transcript"], transcription.text),
        ]:
            if not value:
                continue
            y -= 8
            draw_line(label, 12, colors.HexColor("#0f172a"), 20)
            for raw_line in str(value).splitlines() or [""]:
                for line in textwrap.wrap(raw_line, width=92, replace_whitespace=False) or [""]:
                    draw_line(line, 10, colors.HexColor("#1f2937"), 14)

        pdf.save()
        buffer.seek(0)
        return buffer

    @staticmethod
    def docx(transcription, lang="en"):
        from docx import Document

        label = labels(lang)
        document = Document()
        document.add_heading(transcription.title or label["document_title"], level=1)
        document.add_paragraph(
            f"{label['date']}: {transcription.created_at.strftime('%Y-%m-%d %H:%M') if transcription.created_at else '-'}"
        )
        document.add_paragraph(f"{label['language']}: {transcription.language or '-'}")
        document.add_paragraph(f"{label['words']}: {transcription.word_count or 0}")

        if transcription.summary:
            document.add_heading(label["summary"], level=2)
            document.add_paragraph(transcription.summary)

        if transcription.keywords_json:
            document.add_heading(label["keywords"], level=2)
            document.add_paragraph(", ".join(transcription.keywords_json))

        document.add_heading(label["transcript"], level=2)
        document.add_paragraph(transcription.text or "")

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer
